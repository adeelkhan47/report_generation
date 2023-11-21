import re
from os import listdir
from os.path import isfile, join

from docx import Document
from docx import Document as Document_compose
from docx.shared import Pt, RGBColor
from docxcompose.composer import Composer

foldername = "file"
colors = {"black": "0,0,0", "red": "255,0,0", "blue": "0,0,255", "white": "255,255,255"}
filenames = {
    "russian": "sep_rm.txt",
    "telegram": "sep_tg01.txt",
    "mega": "sep_mega02.txt",
    "test": "test.txt",
    "search": "search_terms.txt",
}
output_filenames = {
    "russian": "russian_output.docx",
    "telegram": "telegram_output.docx",
    "mega": "mega_output.docx",
    "combined": "combined.docx",
}
standard_data_filenames = {
    "russian": "russian.docx",
    "telegram": "telegram.docx",
    "mega": "mega.docx",
}


def _set_cell_background(cell, fill, color=None, val=None):
    """
    @fill: Specifies the color to be used for the background
    @color: Specifies the color to be used for any foreground
    pattern specified with the val attribute
    @val: Specifies the pattern to be used to lay the pattern
    color over the background color.
    """
    from docx.oxml.parser import OxmlElement
    from docx.oxml.shared import qn  # feel free to move these out

    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath("w:shd")[
            0
        ]  # in case there's already shading
    except IndexError:
        cell_shading = OxmlElement("w:shd")  # add new w:shd element to it
    if fill:
        cell_shading.set(qn("w:fill"), fill)  # set fill property, respecting namespace
    if color:
        pass  # TODO
    if val:
        pass  # TODO
    cell_properties.append(cell_shading)


def get_file_names(directory_name):
    return [f for f in listdir(directory_name) if isfile(join(directory_name, f))]


def get_file_content(filename):
    with open(filename, "r") as file:
        content = file.read()
        return content


def get_source_actor_n_date(content, filename):
    source, actor, date = None, None, None
    if filename in [filenames["russian"], filenames["mega"]]:
        source = content.split("\n")[0].strip()
    else:
        source = re.search(r"Site: (.+)", content)
        if source:
            source = source.group(1).strip()
    actor = re.search(r"Actor: (.+)", content)
    date = re.search(r"Type: (.+?) \| (.+)", content)
    if actor:
        actor = actor.group(1).strip()
    if date:
        _, date = date.groups()
    else:
        date = re.search(r"Date: (.+)", content).group(1).strip()
    return source, actor, date


def get_market_init_data(content):
    stealer, country, isp, login_count, price = None, None, None, 0, 0
    stealer = re.search(r"Stealer:\s*(\w+)", content)
    country = re.search(r"Country:\s*([\s\S]*?)(?:\n|$)", content)
    isp = re.search(r"ISP:\s*([\s\S]*?)(?:\n\n|$)", content)
    price = re.search(r"Price:([\s\S]*?)(?:\n\n|$)", content)
    login_count = len(re.findall(r"Login: \+", content))

    stealer = stealer.group(1).strip() if stealer else None
    country = " ".join(country.group(1).split("\n")) if country else None
    isp = " ".join(isp.group(1).split()) if isp else None
    price = " ".join(price.group(1).split("\n")) if price else None

    return stealer, country, isp, login_count, price


def get_matching_data(content, domains):
    result = [
        line.strip()
        for line in content.split("\n")
        if any(domain in line for domain in domains)
    ]
    unique_result = list(set(result))
    return unique_result


def get_telegram_init_data(content):
    document_name, document_size = None, None
    document_name = re.search(r"Document Name: (.+)", content)
    document_size = re.search(r"Document Size: (.+)", content)

    document_name = document_name.group(1).strip() if document_name else None
    document_size = document_size.group(1).strip() if document_size else None
    return document_name, document_size


def get_mega_init_data(content):
    document_url = re.search(r"Document URL: (.+)", content)
    document_name = re.search(r"Document Name: (.+)", content)
    document_size = re.search(r"Document Size: (.+)", content)

    document_url = document_url.group(1).strip() if document_url else None
    document_name = document_name.group(1).strip() if document_name else None
    document_size = document_size.group(1).strip() if document_size else None
    return document_url, document_name, document_size


def color_line(key, key_bold, key_color, paragraph):
    key_color = key_color.split(",")
    key_color = [int(color) for color in key_color]
    run_actor_label = paragraph.add_run(key)
    run_actor_label.font.bold = key_bold
    run_actor_label.font.color.rgb = RGBColor(key_color[0], key_color[1], key_color[2])


def add_cell_key_n_color(key, key_color, key_bold, cell_number, cell):
    key_color = key_color.split(",")
    key_color = [int(color) for color in key_color]
    cell_run = cell[cell_number].paragraphs[0].add_run(key)
    cell_run.font.bold = key_bold
    cell_run.font.color.rgb = RGBColor(key_color[0], key_color[1], key_color[2])


def color_table_background(rows, cols, color, table):
    for i in range(rows):
        for j in range(cols):
            _set_cell_background(table.rows[i].cells[j], color)


def add_key_value_in_doc(
    key, value, key_color, value_color, key_bold, value_bold, paragraph
):
    color_line(key, key_bold, key_color, paragraph)
    color_line(value, value_bold, value_color, paragraph)
    paragraph.add_run("\n")


def populate_market_data(doc, content):
    site_sections = content.split("Site:")
    site_sections = [section.strip() for section in site_sections if section.strip()]
    for site in site_sections:
        source, actor, date = get_source_actor_n_date(site, filenames["russian"])
        p = doc.add_paragraph(f"Findings:")
        p.runs[0].font.size = Pt(16)
        p.runs[0].font.bold = True
        rows, cols = 3, 2
        table = doc.add_table(rows, cols)
        source_cells, actor_cells = table.rows[0].cells, table.rows[1].cells
        date_cells = table.rows[2].cells
        add_cell_key_n_color("Source:", colors["white"], False, 0, source_cells)
        add_cell_key_n_color(source, colors["white"], True, 1, source_cells)
        add_cell_key_n_color("Actor:", colors["white"], False, 0, actor_cells)
        add_cell_key_n_color(actor, colors["white"], True, 1, actor_cells)
        add_cell_key_n_color("Crawl Date:", colors["white"], False, 0, date_cells)
        add_cell_key_n_color(date, colors["white"], True, 1, date_cells)
        color_table_background(rows, cols, "#00008B", table)
        start_row, end_row = table.rows[0], table.rows[2]
        start_row.cells[0].merge(end_row.cells[0])
        start_row.cells[1].merge(end_row.cells[1])
        stealer, country, isp, login_count, price = get_market_init_data(site)
        paragraph = doc.add_paragraph()
        add_key_value_in_doc(
            "\tStealer: ",
            stealer,
            colors["black"],
            colors["red"],
            False,
            True,
            paragraph,
        )
        add_key_value_in_doc(
            "\tCountry: ",
            country,
            colors["black"],
            colors["black"],
            False,
            True,
            paragraph,
        )
        add_key_value_in_doc(
            "\tISP: ", isp, colors["black"], colors["black"], False, True, paragraph
        )
        add_key_value_in_doc(
            "\tTotal Resources: ",
            str(login_count),
            colors["black"],
            colors["black"],
            False,
            True,
            paragraph,
        )

        color_line(f"\tResources:", False, colors["black"], paragraph)
        paragraph.add_run("\n")
        resources = get_matching_data(site, domains)
        for resource in resources:
            color_line(f"\t{resource}", False, colors["blue"], paragraph)
            paragraph.add_run("\n")
        add_key_value_in_doc(
            "\tPrice: ", price, colors["black"], colors["black"], False, True, paragraph
        )


def populate_telegram_data(doc, content):
    title_sections = content.split("Title: ")
    title_sections = [section.strip() for section in title_sections if section.strip()]

    for title in title_sections:
        source, actor, date = get_source_actor_n_date(title, filenames["telegram"])
        p = doc.add_paragraph(f"Findings:")
        p.runs[0].font.size = Pt(16)
        p.runs[0].font.bold = True
        rows, cols = 3, 2
        table = doc.add_table(rows, cols)
        source_cells, actor_cells = table.rows[0].cells, table.rows[1].cells
        date_cells = table.rows[2].cells
        add_cell_key_n_color("Source:", colors["white"], False, 0, source_cells)
        add_cell_key_n_color(source, colors["white"], True, 1, source_cells)
        add_cell_key_n_color("Actor:", colors["white"], False, 0, actor_cells)
        add_cell_key_n_color(actor, colors["white"], True, 1, actor_cells)
        add_cell_key_n_color("Crawl Date:", colors["white"], False, 0, date_cells)
        add_cell_key_n_color(date, colors["white"], True, 1, date_cells)
        color_table_background(rows, cols, "#00008B", table)
        start_row, end_row = table.rows[0], table.rows[2]
        start_row.cells[0].merge(end_row.cells[0])
        start_row.cells[1].merge(end_row.cells[1])
        document_name, document_size = get_telegram_init_data(content)
        paragraph = doc.add_paragraph()
        add_key_value_in_doc(
            "\tDocument Name: ",
            document_name,
            colors["black"],
            colors["black"],
            False,
            True,
            paragraph,
        )
        add_key_value_in_doc(
            "\tDocument Size: ",
            document_size,
            colors["black"],
            colors["black"],
            False,
            True,
            paragraph,
        )
        color_line(f"\tDocument Content:", False, colors["black"], paragraph)
        paragraph.add_run("\n")
        document_content = get_matching_data(title, domains)
        for content in document_content:
            name, domain = content.split("@")
            domain, another_name = domain.split(":")
            color_line(f"\t{name}", False, colors["black"], paragraph)
            color_line(f"@{domain}", False, colors["blue"], paragraph)
            color_line(f":{another_name}", False, colors["black"], paragraph)
            paragraph.add_run("\n")


def populate_mega_data(doc, content):
    site_sections = content.split("Site:")
    site_sections = [section.strip() for section in site_sections if section.strip()]
    for site in site_sections:
        source, actor, date = get_source_actor_n_date(site, filenames["mega"])
        p = doc.add_paragraph(f"Findings:")
        p.runs[0].font.size = Pt(16)
        p.runs[0].font.bold = True
        rows, cols = 3, 2
        table = doc.add_table(rows, cols)
        source_cells, actor_cells = table.rows[0].cells, table.rows[1].cells
        date_cells = table.rows[2].cells
        add_cell_key_n_color("Source:", colors["white"], False, 0, source_cells)
        add_cell_key_n_color(source, colors["white"], True, 1, source_cells)
        add_cell_key_n_color("Actor:", colors["white"], False, 0, actor_cells)
        add_cell_key_n_color(actor, colors["white"], True, 1, actor_cells)
        add_cell_key_n_color("Crawl Date:", colors["white"], False, 0, date_cells)
        add_cell_key_n_color(date, colors["white"], True, 1, date_cells)
        color_table_background(rows, cols, "#00008B", table)
        start_row, end_row = table.rows[0], table.rows[2]
        start_row.cells[0].merge(end_row.cells[0])
        start_row.cells[1].merge(end_row.cells[1])
        document_url, document_name, document_size = get_mega_init_data(site)
        paragraph = doc.add_paragraph()
        add_key_value_in_doc(
            "\tDocument URL: ",
            document_url,
            colors["black"],
            colors["black"],
            False,
            True,
            paragraph,
        )
        add_key_value_in_doc(
            "\tDocument Name: ",
            document_name,
            colors["black"],
            colors["black"],
            False,
            True,
            paragraph,
        )
        add_key_value_in_doc(
            "\tDocument Size: ",
            document_size,
            colors["black"],
            colors["black"],
            False,
            True,
            paragraph,
        )
        color_line(f"\tDocument Content:", False, colors["black"], paragraph)
        paragraph.add_run("\n")
        document_content = get_matching_data(site, domains)
        for content in document_content:
            name, domain = content.split("@")
            domain, another_name = domain.split(":")
            color_line(f"\t{name}", False, colors["black"], paragraph)
            color_line(f"@{domain}", False, colors["blue"], paragraph)
            color_line(f":{another_name}", False, colors["black"], paragraph)
            paragraph.add_run("\n")


domains = get_file_content(f"{foldername}/{filenames['search']}")
domains = domains.split("\n")

master = Document_compose(standard_data_filenames["russian"])
composer = Composer(master)
doc = Document()
content = get_file_content(f"{foldername}/{filenames['russian']}")
populate_market_data(doc, content)

doc.save(output_filenames["russian"])
doc2 = Document_compose(output_filenames["russian"])
composer.append(doc2)
composer.save(output_filenames["combined"])

master = Document_compose(output_filenames["combined"])
composer = Composer(master)
doc2 = Document_compose(standard_data_filenames["telegram"])
composer.append(doc2)
composer.save(output_filenames["combined"])

master = Document_compose(output_filenames["combined"])
composer = Composer(master)
doc = Document()
content = get_file_content(f"{foldername}/{filenames['telegram']}")
populate_telegram_data(doc, content)
doc.save(output_filenames["telegram"])
doc2 = Document_compose(output_filenames["telegram"])
composer.append(doc2)
composer.save(output_filenames["combined"])

master = Document_compose(output_filenames["combined"])
composer = Composer(master)
doc2 = Document_compose(standard_data_filenames["mega"])
composer.append(doc2)
composer.save(output_filenames["combined"])

master = Document_compose(output_filenames["combined"])
composer = Composer(master)
doc = Document()
content = get_file_content(f"{foldername}/{filenames['mega']}")
populate_mega_data(doc, content)

doc.save(output_filenames["mega"])
doc2 = Document_compose(output_filenames["mega"])
composer.append(doc2)
composer.save(output_filenames["combined"])
